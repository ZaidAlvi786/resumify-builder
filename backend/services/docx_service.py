from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from schemas.resume import ResumeOutput

def generate_docx_resume(data: ResumeOutput) -> io.BytesIO:
    """
    Generates a professional Word document from resume data.
    """
    doc = Document()
    
    # 1. Header: Name & Contact Info
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_run = header_para.add_run(data.full_name.upper())
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
    
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{data.contact_info.get('email', '')} | {data.contact_info.get('phone', '')} | {data.contact_info.get('location', '')}"
    contact_para.add_run(contact_text).font.size = Pt(10)
    
    # 2. Professional Summary
    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    doc.add_paragraph(data.summary)
    
    # 3. Skills
    doc.add_heading("KEY SKILLS", level=1)
    skills_text = ", ".join(data.skills)
    doc.add_paragraph(skills_text)
    
    # 4. Experience
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    for exp in data.experience:
        exp_para = doc.add_paragraph()
        role_run = exp_para.add_run(f"{exp.title}")
        role_run.bold = True
        exp_para.add_run(f"\t{exp.company} | {exp.start_date} - {exp.end_date}")
        exp_para.paragraph_format.tab_stops.add_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        
        for bullet in exp.bullet_points:
            doc.add_paragraph(bullet, style='List Bullet')
            
    # 5. Education
    doc.add_heading("EDUCATION", level=1)
    for edu in data.education:
        edu_para = doc.add_paragraph()
        edu_run = edu_para.add_run(edu.degree)
        edu_run.bold = True
        edu_para.add_run(f"\t{edu.school}, {edu.graduation_year}")
        edu_para.paragraph_format.tab_stops.add_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        
    # 6. Projects
    if data.projects:
        doc.add_heading("PROJECTS", level=1)
        for proj in data.projects:
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(proj.name)
            proj_run.bold = True
            doc.add_paragraph(proj.description)
            
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
